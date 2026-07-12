"""Extract plain text from uploaded files (PDF / DOCX / MD / TXT).

Used by the web UI to let users drag in a JD in any format instead of pasting.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Tuple


class ExtractError(ValueError):
    """Raised for unsupported types or unparseable files."""


_SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".md", ".markdown", ".txt", ".text",
}


def _extract_pdf(data: bytes) -> str:
    # Local import — pypdf is a real dep but we don't want module-load cost
    # for callers that never see a PDF.
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise ExtractError(f"PDF parse failed: {e}") from e
    pages: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            t = ""
        if t.strip():
            pages.append(t)
    return "\n\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise ExtractError(f"DOCX parse failed: {e}") from e
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Many resume/JD .docx files use tables — pull cell text in row order.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_plain(data: bytes) -> str:
    # Try utf-8 first, fall back to latin-1 (handles most resume PDFs-saved-as-text)
    try:
        return data.decode("utf-8").strip()
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace").strip()


def extract_text(filename: str, data: bytes) -> Tuple[str, str]:
    """Return (extracted_text, kind) for the given file.

    Raises ExtractError on unsupported extensions or parse failures.
    """
    ext = Path(filename).suffix.lower()
    if ext not in _SUPPORTED_EXTENSIONS:
        raise ExtractError(
            f"unsupported file type {ext!r}. Supported: {sorted(_SUPPORTED_EXTENSIONS)}"
        )
    if ext == ".pdf":
        return _extract_pdf(data), "pdf"
    if ext == ".docx":
        return _extract_docx(data), "docx"
    # .md / .markdown / .txt / .text
    return _extract_plain(data), "text"
