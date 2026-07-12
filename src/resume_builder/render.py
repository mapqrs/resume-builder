"""Render a resume to .docx, applying a Template for formatting.

Two modes:
1. Master only — render the full master verbatim (useful for `--no-tailor`).
2. Master + TailoredResume — render filtered/reordered/rewritten content.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from .schema import (
    CoverLetter,
    Education,
    Experience,
    FontSpec,
    Master,
    Project,
    SkillGroup,
    TailoredResume,
    Template,
)


# ---------- Helpers ----------


_LENGTH_RE = re.compile(r"^\s*([\d.]+)\s*(in|cm|mm|pt)?\s*$", re.IGNORECASE)


def _parse_length(s: str):
    m = _LENGTH_RE.match(s)
    if not m:
        raise ValueError(f"Invalid length: {s!r}")
    value = float(m.group(1))
    unit = (m.group(2) or "in").lower()
    if unit == "in":
        return Inches(value)
    if unit == "cm":
        return Inches(value / 2.54)
    if unit == "mm":
        return Inches(value / 25.4)
    if unit == "pt":
        return Pt(value)
    raise ValueError(f"Unknown unit: {unit}")


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    if len(h) != 6:
        raise ValueError(f"Invalid hex color: {hex_str!r}")
    return RGBColor.from_string(h.upper())


def _apply_font(run, spec: FontSpec, color_hex: Optional[str] = None) -> None:
    run.font.name = spec.name
    run.font.size = Pt(spec.size)
    run.font.bold = spec.bold
    run.font.italic = spec.italic
    if color_hex:
        run.font.color.rgb = _hex_to_rgb(color_hex)
    # Force the Asian font name too — python-docx workaround so font sticks in Word.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), spec.name)
    rFonts.set(qn("w:ascii"), spec.name)
    rFonts.set(qn("w:hAnsi"), spec.name)


def _apply_paragraph_spacing(para, template: Template) -> None:
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = template.spacing.line
    pf.space_after = Pt(template.spacing.paragraph_after)
    pf.space_before = Pt(0)


def _apply_page(doc: _Document, template: Template) -> None:
    section = doc.sections[0]
    section.top_margin = _parse_length(template.page.margin_top)
    section.bottom_margin = _parse_length(template.page.margin_bottom)
    section.left_margin = _parse_length(template.page.margin_left)
    section.right_margin = _parse_length(template.page.margin_right)
    if template.page.size == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    elif template.page.size == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)


def _contact_line(master: Master, *, include_links: bool = True) -> str:
    contact_bits = []
    if master.basics.email:
        contact_bits.append(master.basics.email)
    if master.basics.phone:
        contact_bits.append(master.basics.phone)
    if master.basics.location:
        contact_bits.append(master.basics.location)
    if include_links:
        for link in master.basics.links:
            contact_bits.append(f"{link.label}: {link.url}")
    return " | ".join(contact_bits)


def _add_page_number_run(para) -> None:
    run = para.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def _render_footer(doc: _Document, master: Master, template: Template) -> None:
    """Add a compact footer so separated pages retain contact context."""
    section = doc.sections[0]
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_paragraph_spacing(para, template)

    contact = _contact_line(master, include_links=False)
    text = master.basics.name
    if contact:
        text += f" | {contact}"
    text += " | Page "
    run = para.add_run(text)
    _apply_font(
        run,
        FontSpec(name=template.fonts.body.name, size=max(template.fonts.body.size - 1, 8)),
        template.colors.body,
    )
    _add_page_number_run(para)


# ---------- Tailored projection ----------


def _project_master(
    master: Master, tailored: Optional[TailoredResume]
) -> tuple[Optional[str], list[Experience], list[Project]]:
    """Return (summary, experience_list, project_list) after applying tailored projection.

    If tailored is None, returns the master as-is.
    """
    if tailored is None:
        return master.summary, master.experience, master.projects

    summary = tailored.summary or master.summary

    # Build a lookup of tailored items per section
    exp_items: list[tuple[str, list[tuple[str, str]]]] = []  # (container_id, [(bullet_id, text)])
    proj_items: list[tuple[str, list[tuple[str, str]]]] = []

    for section in tailored.sections:
        if section.name == "experience":
            for item in section.items:
                exp_items.append(
                    (item.source_id, [(b.source_id, b.rewritten_text) for b in item.bullets])
                )
        elif section.name == "projects":
            for item in section.items:
                proj_items.append(
                    (item.source_id, [(b.source_id, b.rewritten_text) for b in item.bullets])
                )

    projected_exp: list[Experience] = []
    for container_id, bullets in exp_items:
        src = master.container_by_id(container_id)
        if not isinstance(src, Experience):
            continue
        new_bullets = []
        for bullet_id, text in bullets:
            orig = master.bullet_by_id(bullet_id)
            if not orig:
                continue
            # Replace text with rewritten version, keep id and tags
            new_bullets.append(orig.model_copy(update={"text": text}))
        projected_exp.append(src.model_copy(update={"bullets": new_bullets}))

    projected_proj: list[Project] = []
    for container_id, bullets in proj_items:
        src = master.container_by_id(container_id)
        if not isinstance(src, Project):
            continue
        new_bullets = []
        for bullet_id, text in bullets:
            orig = master.bullet_by_id(bullet_id)
            if not orig:
                continue
            new_bullets.append(orig.model_copy(update={"text": text}))
        projected_proj.append(src.model_copy(update={"bullets": new_bullets}))

    return summary, projected_exp, projected_proj


# ---------- Section renderers ----------


def _render_header(doc: _Document, master: Master, template: Template) -> None:
    para = doc.add_paragraph()
    _apply_paragraph_spacing(para, template)
    name_run = para.add_run(master.basics.name)
    _apply_font(name_run, template.fonts.name, template.colors.heading)

    contact = _contact_line(master)
    if contact:
        contact_para = doc.add_paragraph()
        _apply_paragraph_spacing(contact_para, template)
        contact_run = contact_para.add_run(contact)
        _apply_font(contact_run, template.fonts.body, template.colors.body)


def _render_section_heading(doc: _Document, title: str, template: Template) -> None:
    para = doc.add_paragraph()
    _apply_paragraph_spacing(para, template)
    run = para.add_run(title.upper())
    _apply_font(run, template.fonts.heading, template.colors.accent)
    # Border below heading
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), template.colors.accent.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_summary(doc: _Document, summary: Optional[str], template: Template) -> None:
    if not summary:
        return
    _render_section_heading(doc, "Summary", template)
    para = doc.add_paragraph()
    _apply_paragraph_spacing(para, template)
    run = para.add_run(summary.strip())
    _apply_font(run, template.fonts.body, template.colors.body)


def _render_experience(
    doc: _Document, items: Iterable[Experience], template: Template
) -> None:
    items = list(items)
    if not items:
        return
    _render_section_heading(doc, "Experience", template)
    for exp in items:
        # Role @ Company line
        role_para = doc.add_paragraph()
        _apply_paragraph_spacing(role_para, template)
        role_run = role_para.add_run(f"{exp.role} — {exp.company}")
        _apply_font(role_run, template.fonts.role, template.colors.heading)

        # Dates / location line
        meta_bits = [f"{exp.start} – {exp.end}"]
        if exp.location:
            meta_bits.append(exp.location)
        meta_para = doc.add_paragraph()
        _apply_paragraph_spacing(meta_para, template)
        meta_run = meta_para.add_run(" | ".join(meta_bits))
        _apply_font(
            meta_run,
            FontSpec(name=template.fonts.body.name, size=template.fonts.body.size, italic=True),
            template.colors.body,
        )

        for bullet in exp.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            _apply_paragraph_spacing(bp, template)
            run = bp.add_run(bullet.text)
            _apply_font(run, template.fonts.body, template.colors.body)


def _render_projects(
    doc: _Document, items: Iterable[Project], template: Template
) -> None:
    items = list(items)
    if not items:
        return
    _render_section_heading(doc, "Projects", template)
    for proj in items:
        head_para = doc.add_paragraph()
        _apply_paragraph_spacing(head_para, template)
        head_run = head_para.add_run(proj.name)
        _apply_font(head_run, template.fonts.role, template.colors.heading)
        if proj.url:
            url_run = head_para.add_run(f" — {proj.url}")
            _apply_font(url_run, template.fonts.body, template.colors.body)

        for bullet in proj.bullets:
            bp = doc.add_paragraph(style="List Bullet")
            _apply_paragraph_spacing(bp, template)
            run = bp.add_run(bullet.text)
            _apply_font(run, template.fonts.body, template.colors.body)


def _education_line(edu: Education) -> str:
    """Status-aware first line for an education entry.

    Different statuses need visibly different phrasing on the resume —
    a dropout shouldn't read as a graduate, a certification shouldn't
    read as a degree, a deferred admit shouldn't read as an enrolment.
    """
    status = edu.status
    if status == "in_progress":
        line = f"{edu.degree} (in progress), {edu.school} (expected {edu.year})"
    elif status == "dropout":
        line = f"Attended {edu.school} — {edu.degree} ({edu.year})"
    elif status == "deferred_admit":
        line = f"Admitted to {edu.school} — {edu.degree} (deferred from {edu.year})"
    elif status == "rejected_admit":
        line = f"Offered admission to {edu.school} — {edu.degree} (declined, {edu.year})"
    elif status == "on_leave":
        line = f"{edu.degree} (on leave), {edu.school} ({edu.year})"
    elif status == "certification_only":
        line = f"{edu.degree}, {edu.school} ({edu.year})"
    elif status == "online_only":
        line = f"{edu.degree} (online), {edu.school} ({edu.year})"
    else:  # graduated (default)
        line = f"{edu.degree}, {edu.school} ({edu.year})"
    if edu.location:
        line += f" — {edu.location}"
    if edu.gpa:
        line += f" · {edu.gpa}"
    if edu.reason:
        # Position-of-strength narrative — user-supplied, free-form.
        # Surfaces what the status alone can't (e.g. dropout WHY).
        line += f"; {edu.reason}"
    return line


def _render_education(
    doc: _Document, items: Iterable[Education], template: Template
) -> None:
    items = list(items)
    if not items:
        return
    _render_section_heading(doc, "Education", template)
    for edu in items:
        para = doc.add_paragraph()
        _apply_paragraph_spacing(para, template)
        run = para.add_run(_education_line(edu))
        _apply_font(run, template.fonts.body, template.colors.body)

        # Notes (italic, one paragraph).
        if edu.notes:
            note_para = doc.add_paragraph()
            _apply_paragraph_spacing(note_para, template)
            note_run = note_para.add_run(edu.notes)
            _apply_font(
                note_run,
                FontSpec(name=template.fonts.body.name, size=template.fonts.body.size, italic=True),
                template.colors.body,
            )

        # Awards (one paragraph each, indented). Bock requires criteria —
        # awards without it still render, but with a "(no criteria)" hint
        # so the user notices and adds context.
        for award in edu.awards:
            award_para = doc.add_paragraph()
            _apply_paragraph_spacing(award_para, template)
            parts = [award.name]
            if award.criteria:
                parts.append(award.criteria)
            else:
                parts.append("(no criteria listed)")
            if award.year:
                parts.append(award.year)
            award_run = award_para.add_run("— " + " · ".join(parts))
            _apply_font(award_run, template.fonts.body, template.colors.body)


def _render_skills(
    doc: _Document, items: Iterable[SkillGroup], template: Template
) -> None:
    items = list(items)
    if not items:
        return
    _render_section_heading(doc, "Skills", template)
    for group in items:
        para = doc.add_paragraph()
        _apply_paragraph_spacing(para, template)
        cat_run = para.add_run(f"{group.category}: ")
        _apply_font(
            cat_run,
            FontSpec(
                name=template.fonts.body.name,
                size=template.fonts.body.size,
                bold=True,
            ),
            template.colors.body,
        )
        items_run = para.add_run(", ".join(group.items))
        _apply_font(items_run, template.fonts.body, template.colors.body)


# ---------- Public entry point ----------


def render_docx(
    master: Master,
    template: Template,
    out_path: str | Path,
    tailored: Optional[TailoredResume] = None,
) -> Path:
    """Render the resume to a .docx file. Returns the output path."""
    doc = Document()
    _apply_page(doc, template)
    _render_footer(doc, master, template)

    summary, experience, projects = _project_master(master, tailored)

    _render_header(doc, master, template)

    for section_name in template.section_order:
        if section_name == "summary":
            _render_summary(doc, summary, template)
        elif section_name == "experience":
            _render_experience(doc, experience, template)
        elif section_name == "projects":
            _render_projects(doc, projects, template)
        elif section_name == "education":
            _render_education(doc, master.education, template)
        elif section_name == "skills":
            _render_skills(doc, master.skills, template)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


# ---------- Cover letter ----------


def render_cover_letter_docx(
    master: Master,
    template: Template,
    cover: CoverLetter,
    out_path: str | Path,
) -> Path:
    """Render a CoverLetter to .docx using the same template fonts/margins as the resume.

    Layout (top-to-bottom):
        Name (heading font)
        Contact line (body font, muted)
        Links line (body font)
        <blank>
        Salutation
        <blank>
        Paragraph 1
        Paragraph 2
        ...
        Closing,
        Name
    """
    from datetime import date

    doc = Document()
    _apply_page(doc, template)

    # Letterhead — same as resume header style
    name_para = doc.add_paragraph()
    _apply_paragraph_spacing(name_para, template)
    name_run = name_para.add_run(master.basics.name)
    _apply_font(name_run, template.fonts.name, color_hex=template.colors.heading)

    contact_bits = [s for s in (master.basics.email, master.basics.phone, master.basics.location) if s]
    if contact_bits:
        contact_para = doc.add_paragraph()
        _apply_paragraph_spacing(contact_para, template)
        c_run = contact_para.add_run(" · ".join(contact_bits))
        _apply_font(c_run, template.fonts.body)

    if master.basics.links:
        link_para = doc.add_paragraph()
        _apply_paragraph_spacing(link_para, template)
        link_text = " · ".join(f"{lk.label}: {lk.url}" for lk in master.basics.links)
        l_run = link_para.add_run(link_text)
        _apply_font(l_run, template.fonts.body, color_hex=template.colors.accent)

    # Date
    date_para = doc.add_paragraph()
    _apply_paragraph_spacing(date_para, template)
    d_run = date_para.add_run(date.today().strftime("%B %-d, %Y") if hasattr(date.today(), "strftime") else str(date.today()))
    _apply_font(d_run, template.fonts.body)

    # Salutation
    sal_para = doc.add_paragraph()
    _apply_paragraph_spacing(sal_para, template)
    s_run = sal_para.add_run(cover.salutation)
    _apply_font(s_run, template.fonts.body)

    # Body paragraphs
    for p in cover.paragraphs:
        body_para = doc.add_paragraph()
        _apply_paragraph_spacing(body_para, template)
        b_run = body_para.add_run(p.text.strip())
        _apply_font(b_run, template.fonts.body)

    # Closing + signature
    close_para = doc.add_paragraph()
    _apply_paragraph_spacing(close_para, template)
    cl_run = close_para.add_run(cover.closing)
    _apply_font(cl_run, template.fonts.body)

    sig_para = doc.add_paragraph()
    _apply_paragraph_spacing(sig_para, template)
    sg_run = sig_para.add_run(master.basics.name)
    _apply_font(sg_run, template.fonts.body)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out
