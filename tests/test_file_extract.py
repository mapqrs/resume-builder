"""Tests for file_extract.py — plain text from PDF / DOCX / MD / TXT."""

from __future__ import annotations

import io

import pytest

from resume_builder.file_extract import ExtractError, extract_text


def test_txt_utf8_roundtrip():
    text = "Senior Backend Engineer\n\nRequirements:\n- 5+ years Python\n"
    out, kind = extract_text("jd.txt", text.encode("utf-8"))
    assert out == text.strip()
    assert kind == "text"


def test_md_treated_as_text():
    md = "# Senior Engineer\n\n**Requirements**\n- Python\n- Postgres\n"
    out, kind = extract_text("role.md", md.encode("utf-8"))
    assert "Senior Engineer" in out
    assert "Postgres" in out
    assert kind == "text"


def test_latin1_fallback():
    """A file with non-utf8 bytes shouldn't crash — fall back to latin-1."""
    bytes_with_latin1 = "Naïve résumé".encode("latin-1")
    out, kind = extract_text("role.txt", bytes_with_latin1)
    assert "Na" in out
    assert "r" in out
    assert kind == "text"


def test_unsupported_extension_rejected():
    with pytest.raises(ExtractError, match="unsupported"):
        extract_text("weird.exe", b"\x00\x01\x02")
    with pytest.raises(ExtractError):
        extract_text("photo.jpg", b"\xff\xd8\xff")


def test_no_extension_rejected():
    with pytest.raises(ExtractError, match="unsupported"):
        extract_text("README", b"hello")


def test_docx_roundtrip_paragraphs():
    """Generate a real .docx, extract, verify content survives."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Senior Backend Engineer")
    doc.add_paragraph("")
    doc.add_paragraph("Requirements:")
    doc.add_paragraph("5+ years Python")
    doc.add_paragraph("Deep Postgres experience")

    buf = io.BytesIO()
    doc.save(buf)

    out, kind = extract_text("jd.docx", buf.getvalue())
    assert kind == "docx"
    assert "Senior Backend Engineer" in out
    assert "5+ years Python" in out
    assert "Postgres" in out


def test_docx_table_cells_included():
    """Resume / JD .docx files often use tables; cell text should come through."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Skills")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Languages"
    table.rows[0].cells[1].text = "Python, Go"
    table.rows[1].cells[0].text = "Infra"
    table.rows[1].cells[1].text = "Kubernetes, Postgres"

    buf = io.BytesIO()
    doc.save(buf)

    out, _ = extract_text("resume.docx", buf.getvalue())
    assert "Languages" in out
    assert "Python, Go" in out
    assert "Kubernetes, Postgres" in out


def test_pdf_roundtrip():
    """Generate a tiny one-page PDF via pypdf, extract, verify content."""
    # pypdf can WRITE PDFs too — use that to avoid a binary fixture in the repo.
    pytest.importorskip("pypdf")
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DecodedStreamObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    # Build a minimal valid PDF with one page containing visible text.
    # We bypass pypdf's high-level helpers (which require fonts) and write
    # a hand-rolled content stream with a built-in font.
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    # Built-in Helvetica font
    font = DictionaryObject({
        NameObject("/Type"): NameObject("/Font"),
        NameObject("/Subtype"): NameObject("/Type1"),
        NameObject("/BaseFont"): NameObject("/Helvetica"),
    })
    font_ref = writer._add_object(font)
    resources = DictionaryObject({
        NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref}),
    })
    page[NameObject("/Resources")] = resources

    # Content stream: place "Senior Backend Engineer" and "Python Postgres"
    content = (
        b"BT /F1 14 Tf 72 720 Td (Senior Backend Engineer) Tj ET\n"
        b"BT /F1 12 Tf 72 690 Td (Python Postgres Kubernetes) Tj ET\n"
    )
    stream = DecodedStreamObject()
    stream.set_data(content)
    page[NameObject("/Contents")] = writer._add_object(stream)

    buf = io.BytesIO()
    writer.write(buf)

    out, kind = extract_text("jd.pdf", buf.getvalue())
    assert kind == "pdf"
    assert "Senior Backend Engineer" in out
    assert "Python" in out
    assert "Postgres" in out


def test_corrupt_pdf_raises_extracterror():
    with pytest.raises(ExtractError, match="PDF parse failed"):
        extract_text("broken.pdf", b"not a real pdf")


def test_corrupt_docx_raises_extracterror():
    with pytest.raises(ExtractError, match="DOCX parse failed"):
        extract_text("broken.docx", b"not a real docx")
